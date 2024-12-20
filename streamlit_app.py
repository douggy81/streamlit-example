import os
import streamlit as st
from pinecone import Pinecone
from llama_index.vector_stores.pinecone import PineconeVectorStore
from llama_index.core import VectorStoreIndex, Settings, ServiceContext
from llama_index.llms.gemini import Gemini
from llama_index.embeddings.gemini import GeminiEmbedding
from llama_index.core.indices.postprocessor import SentenceEmbeddingOptimizer
import google.generativeai as genai
from llama_index.core.callbacks import LlamaDebugHandler, CallbackManager
from google.ai.generativelanguage import (
    GenerateAnswerRequest,
    HarmCategory,
    SafetySetting,
)

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from io import BytesIO
import markdown
from bs4 import BeautifulSoup # <-- Corrected import
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch

from datetime import datetime

# --- Streamlit App Config ---
st.set_page_config(
    page_title="Chat with the Gemini, your personal trainer in sales using a methodology developped by Patrick Gassier",
    page_icon="",
    layout="centered",
    menu_items=None
)

# --- Custom CSS to inject for setting the background color to a very light orange ---
def set_light_orange_background():
    css_style = """
    <style>
    .stApp {
        background-color: #FFE0B2;  /* Very light orange */
    }
    </style>
    """
    st.markdown(css_style, unsafe_allow_html=True)
set_light_orange_background()

# --- Safety config ---
safety_config = [
    SafetySetting(
        category=HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT,
        threshold=SafetySetting.HarmBlockThreshold.BLOCK_LOW_AND_ABOVE,
    ),
    SafetySetting(
        category=HarmCategory.HARM_CATEGORY_VIOLENCE,
        threshold=SafetySetting.HarmBlockThreshold.BLOCK_LOW_AND_ABOVE,
    ),
    SafetySetting(
        category=HarmCategory.HARM_CATEGORY_HARASSMENT,
        threshold=SafetySetting.HarmBlockThreshold.BLOCK_ONLY_HIGH,
    ),
]

print("***Manuel de Formation à la vente***")

# --- Global settings for LLM & Embeddings ---
Settings.llm = Gemini(model_name="models/gemini-2.0-flash-exp", api_key=os.environ.get("GOOGLE_API_KEY"))
Settings.embed_model = GeminiEmbedding(model_name="models/text-embedding-004", api_key=os.environ.get("GOOGLE_API_KEY"), embed_batch_size=100)    

# --- Callback manager ---
llama_debug=LlamaDebugHandler(print_trace_on_end=True)
callback_manager=CallbackManager(handlers=[llama_debug])
Settings.callback_manager=callback_manager

print(callback_manager)

# --- Index loading (cached) ---
@st.cache_resource(show_spinner=False)
def get_index() -> VectorStoreIndex:
    try:
        pc = Pinecone(api_key=os.environ.get("PINECONE_API_KEY"))
        index_name = "artofsaletestalpha"
        pinecone_index = pc.Index(index_name=index_name, host="https://artofsaletestalpha-udkb0ne.svc.aped-4627-b74a.pinecone.io")
        vector_store = PineconeVectorStore(pinecone_index=pinecone_index)
        return VectorStoreIndex.from_vector_store(vector_store=vector_store)
    except Exception as e:
        st.error(f"Error connecting to Pinecone: {e}")
        return None

index = get_index()


# --- Dynamic Title ---
def update_title():
    title_text = "Chat with the Gemini, your personal trainer in sales using a methodology developed by Patrick Gassier" if st.session_state.selected_language == "English" else "Conversation avec votre formateur personnel sur les méthodologies de vente créées par Patrick Gassier"
    st.title(title_text)
    
# Function to generate the initial greeting
def generate_greeting(language):
    llm_prompt = f"""
        You are a chatbot and a trainer on the book "Manuel de Formation a la vente" (The Art of Sale). 
        You speak English and French. Greet the user and briefly introduce yourself and your capabilities.
        Make sure to reply in the selected language: {language}
        Also, ask for specific details because you are here to help the user on their specific case.
        For example, you could mention that you can simulate a real conversation on the phone for role play and help them practice real life situation.
        The other important fact is that you need to ask for the password which you know and is taxi.
        Unless the user enters it, you won't be able to continue so keep asking for that password until the user enters it, then you can continue.
        Thanks and have fun training a world of eager learners on how to sale and grow their business! 
        One last thing, you cannot give out any clues about the password. If the user doesn't know it, you can't give them clues.
        Feel free to use emojis (like 👍 or 😀) in your responses to make them more engaging.
        """
    try:
        response = st.session_state.chat_engine.chat(message=llm_prompt)
        return response.response
    except Exception as e:
        st.error(f"Error generating greeting: {e}")
        return "Sorry, I am having trouble generating the initial greeting."

# --- Initialize chat engine ---
if "chat_engine" not in st.session_state:
    if index:  # Only initialize if the index loaded properly
        postprocessor = SentenceEmbeddingOptimizer(embed_model=Settings.embed_model, percentile_cutoff=0, threshold_cutoff=0)
        st.session_state.chat_engine = index.as_chat_engine(
            chat_mode="context",
            verbose=True,
            node_postprocessors=[postprocessor],
            system_prompt=f"""
            You are a chatbot and a trainer on this book called the art of sale book, named in french "L'Art de la Vente - Une méthode à la française", you speak English and French.
            This book is written by Patrick Gassier, he was born in Marseille in 1949.
            You will need to greet the user right away and figure out at first a quick intro on what your capable of and what you can help the user with.
            Also, ask for specific details because you are here to help the user on their specific case.
            For example, you could mention that you can simulate a real conversation on the phone for role play and help them practice real life situation.
            The other important fact is that you need to ask for the password which you know and is taxi.
            Unless the user enters it, you won't be able to continue so keep asking for that password until the user enters it, then you can continue.
            Thanks and have fun training a world of eager learners on how to sale and grow their business! 
            One last thing, you cannot give out any clues about the password. If the user doesn't know it, you can't give them clues.
            As soon as the user enters the correct password, greet them by saying that you're very excited to be able to help him/her out and congratulate them as well for their purchase and interest in growing their business.
            That you are here to help. Give a few examples on what you can do related to the book.
            Again, you are a chatbot and trainer, you can generate quizzes as well but mostly you are here to help the user reach their goal to grow their business, feel that they are learning from you great things.
            Don't hesitate to use the best of your knowledge and practice as well as giving out references in the correct language.
            Also, tell the user that if they need a print out of the output, just select Export to Word button directly from the app and format your answer in a nice and printable format.
            Give them example of responses that can be printed such as a checklist for a pre-meeting with a customer that can be printed out and taken on the road.
            Feel free to use emojis (like 👍 or 😀) in your responses to make them more engaging and your tone should be enthousiastic and professional, don't hesistate to be very detailed as well as adding references not only to the book content.
            You also need to make references to best practices of the industry.
            When responding to the user, be sure to use the same language as their message.
            """
        )
    else:
        st.error("Could not initialize the chat engine due to Pinecone connection failure.")


# --- Initialize messages list if not in session state ---
if "messages" not in st.session_state:
    st.session_state.messages = []

# --- Language selection logic ---
if 'selected_language' not in st.session_state:
    st.session_state.selected_language = "Français"  # Default language
update_title()
temp_language = st.selectbox(
    label="Choose your language / Choisissez votre langue",
    options=["English", "Français"],
    index=["English", "Français"].index(st.session_state.selected_language),
    key="language_select"
)
confirm_button = st.button(label="Confirm / Confirmer")

if confirm_button:
    st.session_state.selected_language = temp_language
    st.session_state.messages = []  # Clear previous messages
    with st.spinner("Generating greeting..." if st.session_state.selected_language == "English" else "Génération du message de bienvenue à l'utilisateur..."):
        greeting = generate_greeting(st.session_state.selected_language)
        st.session_state.messages.append({"role": "assistant", "content": greeting})
    update_title()


# --- Chat interface ---
chat_text="Your question..." if st.session_state.selected_language == "English" else "Votre question..."
prompt = st.chat_input(chat_text)

if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.write(message["content"])

# --- Display book link based on language ---
book_link_en = "📖 To get a copy of the book, please click on the following link: [The Art of Sale - A French Method](https://amzn.eu/d/04FT23KE) (The french version is now available on Amazon.fr. The english version is in the works)"
book_link_fr = "📖 Pour obtenir une copie du livre, cliquez sur le lien suivant : [L'Art de la Vente - Une méthode à la française](https://amzn.eu/d/04FT23KE) (la version française est disponible maintenant sur Amazon.fr)"

if st.session_state.selected_language == "English":
    st.markdown(book_link_en)
else:
    st.markdown(book_link_fr)


# --- Helper Functions ---

def format_chat_history(messages):
    """Formats the chat history into a string suitable for document export."""
    formatted_text = ""
    for message in messages:
        formatted_text += f"{message['role'].capitalize()}: {message['content']}\n\n"
    return formatted_text

def create_word_document(formatted_text):
    """Creates a Word document in memory from the formatted text with header."""
    document = Document()

    # Add header with page number and report title
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Date for title
    now = datetime.now()
    formatted_date = now.strftime("%Y-%m-%d")
    report_title = f"Report Generated on {formatted_date} by AI trained on Patrick Gassier's Book 'The Art of Sale'"
    
    run = header_paragraph.add_run(report_title)
    run.bold = True
    header_paragraph.add_run("\t\t\t") # Add tab for spacing

    
    # Add page number logic
    run = header_paragraph.add_run("Page ")
    run.italic = True
    run = header_paragraph.add_run(" \t") # space after "page "
    run.font.superscript = True

    run = header_paragraph.add_run("PAGE ")
    run.font.superscript = True

    run = header_paragraph.add_run("  \t") # tab

    run = header_paragraph.add_run("of ")
    run.font.superscript = True
    run.italic = True
    run = header_paragraph.add_run("  \t")
    run.font.superscript = True

    run = header_paragraph.add_run("NUMPAGES") # Inserted Field 
    run.font.superscript = True
    
    # Add the chat history
    for line in formatted_text.split('\n\n'):
        if line.strip():  # Check if the line has content after removing whitespace
            p = document.add_paragraph()
            html = markdown.markdown(line)
            try: # add the try/except around the soup creation
                soup = BeautifulSoup(html, 'html.parser')

                if soup.body:
                    if soup.body.contents:
                        for element in soup.body.contents:
                            if element.name == 'p':
                                for item in element.contents:
                                    if str(item).startswith('<strong>'):
                                        p.add_run(item.text).bold = True
                                    elif str(item).startswith('<em>'):
                                        p.add_run(item.text).italic = True
                                    else:
                                        p.add_run(str(item))
                    else:
                        document.add_paragraph(line)  # Add the original line if no <p> tags or no body
                else: # Add the original line if there is no body
                     document.add_paragraph(line)
            except AttributeError:  # Handle very rare cases of malformed HTML or if soup is None, etc
                document.add_paragraph(line)
        else:
            document.add_paragraph() # Add empty paragraph for spacing

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# --- Handle LLM responses ---
if st.session_state.messages and st.session_state.messages[-1]["role"] != "assistant":
    with st.chat_message("assistant"):
        spinner_text = "Processing..." if st.session_state.selected_language == "English" else "En cours de traitement..."
        with st.spinner(spinner_text):
            try:
                response = st.session_state.chat_engine.chat(message=prompt)
                raw_response = response.response
                if "__TASK_START__" in raw_response and "__TASK_END__" in raw_response:
                   task_section=raw_response.split("__TASK_START__")[1].split("__TASK_END__")[0]
                   tasks=task_section.strip().split("\n")
                   st.markdown("## Checklist")
                   for i, task in enumerate(tasks):
                      st.checkbox(task.strip(), key=f"task_{i}")
                   response_output = raw_response.split("__TASK_END__")[1].strip()
                   st.write(response_output)
                else:
                    st.write(raw_response) # default output

                message = {"role": "assistant", "content": response.response}
                st.session_state.messages.append(message)
            except Exception as e:
                st.error(f"An error occurred while processing the response: {e}")

# --- Export Functionality ---
# Multilingual button labels
word_label_en = "Export to Word (.docx)"
word_label_fr = "Exporter au format Word (.docx)"

# Initialize word_label before using them
word_label = word_label_en # assign the english default

if st.session_state.selected_language == "English":
    word_label = word_label_en
else:
    word_label = word_label_fr

if st.session_state.messages:
    formatted_history = format_chat_history(st.session_state.messages)
        
    # Word export
    word_buffer = create_word_document(formatted_history)
    st.download_button(
        label=word_label,
        data=word_buffer,
        file_name="chat_history.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
